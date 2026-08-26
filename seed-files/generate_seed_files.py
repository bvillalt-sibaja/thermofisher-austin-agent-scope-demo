#!/usr/bin/env python3
"""Generates the three seed .xlsx workbooks + the Word seed doc for the
Thermo Fisher Austin RPA demo. Re-run any time to reset demo data to a
known-good state.
"""
import os
import random
import sys

import openpyxl
from openpyxl.styles import Font, PatternFill

random.seed(42)
HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, os.path.join(HERE, "..", "excel-mirror"))
from workbook import WorkbookModel  # noqa: E402 - needs sys.path set first

FILL_HEADER = PatternFill(start_color="D9EAD3", end_color="D9EAD3", fill_type="solid")


def header_row(ws, headers):
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.fill = FILL_HEADER


# --------------------------------------------------------------- Production Tracker
def build_production_tracker():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Tracker"
    header_row(ws, ["Priority", "MRP Element Data", "P/N#", "W/O#", "Component Value", "Notes"])

    rows = [
        [1, "5697609", "A42362", "WO-88291", "COMP-4471", ""],
        [2, "5697610", "A35989C", "WO-88292", "COMP-4472", "Awaiting batch release"],
        [3, "5697611", "A30112", "WO-88293", "COMP-4473", ""],
        [2, "5697612", "A41207", "WO-88294", "COMP-4474", "Rush order"],
        [4, "5697613", "A29988", "WO-88295", "COMP-4475", ""],
    ]
    extra_pns = [f"A{random.randint(30000,49999)}" for _ in range(15)]
    for i, pn in enumerate(extra_pns, start=6):
        rows.append([
            random.choice([1, 2, 3, 4]),
            str(5697600 + i),
            pn,
            f"WO-{88290+i}",
            f"COMP-{4470+i}",
            "",
        ])
    for r in rows:
        ws.append(r)

    for col, width in zip("ABCDEF", [9, 16, 10, 10, 14, 34]):
        ws.column_dimensions[col].width = width

    path = os.path.join(HERE, "Production Tracker 2026.xlsx")
    wb.save(path)

    # Real cell formatting (fill/border/merge), matching the recording's
    # "Automatic Color" / "Fill Color RGB in Description" / cosmetic-formatting
    # steps — applied to color/border properties only, never to a cell's
    # position or value, so it can't disturb the row math the automation
    # relies on (find_row_by_value / the first-empty-row fallback).
    m = WorkbookModel(path)
    last_row = len(rows)  # 0-indexed data rows are 1..last_row (row 0 = header)
    # pale-yellow highlight on the Notes column, echoing "Fill Color RGB in
    # Description" (this sheet has no literal "Description" column — Notes
    # is the closest analog, the free-text annotation field)
    for r in range(0, last_row + 1):
        m.set_fill(r, 5, "FFFF99")
    m.set_border_range(0, 0, last_row, 5, style="thin")
    m.set_border_range(0, 0, last_row, 5, style="thick", outline_only=True)
    # section banner, merged & centered, placed past the used columns (H:I)
    # so it can never collide with a data cell an automation reads/writes
    m.merge_and_center(0, 7, 0, 8, value="PRODUCTION TRACKER — LSG PLANT 070")
    m.save()
    return path


# --------------------------------------------------------------- CSAB
def build_csab():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "CSAB"
    header_row(
        ws,
        ["Line/Archive %", "Group", "Owner", "W/O#", "SKU", "Notes", "WIP QTY",
         "Stock Out Date", "Target Due Date"],
    )
    rows = [
        ["97%", "Filling", "AD", "N/A", "A35989C",
         "4/23: Open PR to consume beads to avoid $$$ scrap", 47, "6/29/2026", "TBD"],
        ["88%", "Packaging", "JM", "WO-88291", "A42362", "", 12, "7/2/2026", "7/10/2026"],
        ["72%", "Filling", "RS", "need", "A30112", "Monitoring lot for date extension", 30, "7/15/2026", "TBD"],
    ]
    for r in rows:
        ws.append(r)

    for col, width in zip("ABCDEFGHI", [13, 12, 8, 10, 12, 44, 10, 15, 15]):
        ws.column_dimensions[col].width = width

    path = os.path.join(HERE, "Customer Service Alert Board 2026.xlsx")
    wb.save(path)

    # Real formatting: amber "at risk" fill on the Group cell of the 97%
    # archive row (matches "Select Color Fill in Group" / "Select Fill Color"
    # in the recording), all-borders + thick outside border around the data
    # block, and a merged & centered section banner past the used columns.
    m = WorkbookModel(path)
    last_row = len(rows)
    m.set_fill(1, 1, "FFC000")  # row 2 (0-indexed 1) = the A35989C / 97% / Filling row
    m.set_border_range(0, 0, last_row, 8, style="thin")
    m.set_border_range(0, 0, last_row, 8, style="thick", outline_only=True)
    m.merge_and_center(0, 10, 0, 11, value="CUSTOMER SERVICE ALERT BOARD")
    m.save()
    return path


# --------------------------------------------------------------- V2 Trending Safety Stock
def build_safety_stock():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Safety Stock"
    header_row(ws, ["Finish Good SKU", "Safety Stock Qty", "Trend"])
    rows = [
        ["A35989C", 500, "Stable"],
        ["A42362", 300, "Increasing"],
        ["A30112", 150, "Decreasing"],
    ]
    for r in rows:
        ws.append(r)
    for col, width in zip("ABC", [16, 16, 12]):
        ws.column_dimensions[col].width = width

    path = os.path.join(HERE, "V2 Trending Safety Stock Metric GSD & BID.xlsx")
    wb.save(path)
    return path


# --------------------------------------------------------------- Word doc
def build_word_doc():
    from docx import Document

    doc = Document()
    doc.add_heading("Material Specification Document", level=1)
    doc.add_paragraph("Material: A35989C — Beads, Filling Grade")
    doc.add_paragraph("Plant: 070 — LSG Production")
    doc.add_paragraph(
        "This document is linked to the material master record (Document Data / "
        "Document tab) and describes packaging, storage, and handling "
        "specifications referenced during production order review."
    )
    doc.add_heading("Storage Conditions", level=2)
    doc.add_paragraph("Store between 15-25°C. Protect from moisture.")
    doc.add_heading("Batch Notes", level=2)
    doc.add_paragraph("Lot volume observed: 5.1 L. Date-extended per QA approval.")

    path = os.path.join(HERE, "Material Specification Document.docx")
    doc.save(path)
    return path


if __name__ == "__main__":
    p1 = build_production_tracker()
    p2 = build_csab()
    p3 = build_safety_stock()
    print("Wrote:", p1)
    print("Wrote:", p2)
    print("Wrote:", p3)
    try:
        p4 = build_word_doc()
        print("Wrote:", p4)
    except ImportError:
        print("python-docx not available in this interpreter — run with ~/rpa-env/bin/python3")
